"""Generación automática de documentos Word (docxtpl) de un Préstamo:
solicitud de crédito (expediente KYC del cliente + datos del préstamo),
recibo de desembolso, pagaré y recibo de pago de cuota."""
from io import BytesIO
from pathlib import Path

from docx import Document as DocumentoWord
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from docxtpl import DocxTemplate

from financiera.models import Cliente, Prestamo
from financiera.services.numero_a_letras import monto_en_letras

TEXTO_FIRMA_SOLICITANTE = 'FIRMA DEL SOLICITANTE'
TEXTO_FIRMA_SUSCRIPTOR = 'FIRMA DEL SUSCRIPTOR'
ALTO_FIRMA = Cm(2.2)

# Ciudad de la sede de la institución: no depende del cliente, se usa tanto
# en la Solicitud de Crédito como en el Pagaré.
CIUDAD_INSTITUCION = 'El Rosario, Comayagua'

PLANTILLA_SOLICITUD_PATH = (
    Path(__file__).resolve().parent.parent / 'templates_docx' / 'PLANTILLA_SOLICITUD_DE_CREDITO.docx'
)
PLANTILLA_RECIBO_DESEMBOLSO_PATH = (
    Path(__file__).resolve().parent.parent / 'templates_docx' / 'RECIBO_DE_DESEMBOLSO.docx'
)
PLANTILLA_PAGARE_PATH = Path(__file__).resolve().parent.parent / 'templates_docx' / 'PAGARÉ.docx'
PLANTILLA_RECIBO_PAGO_PATH = Path(__file__).resolve().parent.parent / 'templates_docx' / 'RECIBO_DE_PAGO.docx'

_MESES_ES = {
    1: 'Enero', 2: 'Febrero', 3: 'Marzo', 4: 'Abril', 5: 'Mayo', 6: 'Junio',
    7: 'Julio', 8: 'Agosto', 9: 'Septiembre', 10: 'Octubre', 11: 'Noviembre', 12: 'Diciembre',
}

# Mapas "valor del modelo -> nombre del checkbox en la plantilla de solicitud".
_DESTINO_A_CHK = {
    Prestamo.DESTINO_MEJORAS: 'chk_destino_mejoras',
    Prestamo.DESTINO_PAGO_DEUDA: 'chk_destino_pago_deuda',
    Prestamo.DESTINO_COLEGIATURA: 'chk_destino_colegiatura',
    Prestamo.DESTINO_GASTOS_MEDICOS: 'chk_destino_gastos_medicos',
    Prestamo.DESTINO_CONSUMO: 'chk_destino_consumo',
}
_TIPO_ID_A_CHK = {
    Cliente.TIPO_ID_IDENTIDAD: 'chk_id_identidad',
    Cliente.TIPO_ID_CARNET_RESIDENCIA: 'chk_id_residencia',
    Cliente.TIPO_ID_PASAPORTE: 'chk_id_pasaporte',
}
_GENERO_A_CHK = {
    Cliente.GENERO_FEMENINO: 'chk_genero_femenino',
    Cliente.GENERO_MASCULINO: 'chk_genero_masculino',
}
_ESTADO_CIVIL_A_CHK = {
    Cliente.ESTADO_CIVIL_SOLTERO: 'chk_civil_soltero',
    Cliente.ESTADO_CIVIL_CASADO: 'chk_civil_casado',
    Cliente.ESTADO_CIVIL_DIVORCIADO: 'chk_civil_divorciado',
    Cliente.ESTADO_CIVIL_UNION_LIBRE: 'chk_civil_union_libre',
}
_TIPO_EMPRESA_A_CHK = {
    Cliente.TIPO_EMPRESA_PRIVADA: 'chk_empresa_privada',
    Cliente.TIPO_EMPRESA_PUBLICA: 'chk_empresa_publica',
    Cliente.TIPO_EMPRESA_ONG: 'chk_empresa_ong',
}
_TIPO_EMPLEADO_A_CHK = {
    Cliente.TIPO_EMPLEADO_INDEPENDIENTE: 'chk_empleado_profesional',
    Cliente.TIPO_EMPLEADO_PROPIETARIO: 'chk_empleado_propietario',
    Cliente.TIPO_EMPLEADO_COMERCIANTE: 'chk_empleado_comerciante',
}
_RANGO_SALARIO_A_CHK = {f'RANGO_{i}': f'chk_salario_{i}' for i in range(1, 8)}

# Placeholders que la plantilla oficial pide pero el sistema todavía no
# recolecta (préstamos con el RAP, parentesco con el patrono, PEP, cuentas
# bancarias) -- se dejan siempre sin marcar (casilla vacía).
_CAMPOS_NO_RECOLECTADOS = {
    'chk_rap_ninguno': ' ☐', 'chk_rap_1': ' ☐', 'chk_rap_2': ' ☐',
    'chk_parentesco_ninguno': ' ☐', 'chk_parentesco_1er': ' ☐', 'chk_parentesco_2do': ' ☐', 'chk_parentesco_3er': ' ☐',
    'parentesco_especifique': '',
    'chk_pep_si': ' ☐', 'chk_pep_no': ' ☐',
    'banco_1': '', 'cuenta_1': '', 'banco_2': '', 'cuenta_2': '',
}


def _con_sangria(contexto, espacios=4):
    """Antepone espacios a cada texto ingresado en la Solicitud de Crédito
    (para que se vea mejor formateado, ej. 'Celular:     90909090'). No
    aplica a las casillas 'chk_*' -- esas ya tienen su propio espaciado,
    calibrado para que la casilla no quede pegada a la opción anterior ni
    provoque saltos de línea en medio de una opción."""
    relleno = ' ' * espacios
    resultado = {}
    for clave, valor in contexto.items():
        if clave.startswith('chk_') or valor in (None, ''):
            resultado[clave] = valor
        else:
            resultado[clave] = f'{relleno}{valor}'
    return resultado


def _marcar(mapa_chk, valor_seleccionado):
    """Devuelve {nombre_chk: ' ☒'/' ☐'}, marcando con una casilla visible la
    opción elegida y dejando las demás como casilla vacía. El espacio inicial
    separa la casilla de la etiqueta ANTERIOR (en la plantilla, cada casilla
    queda pegada al final de la opción previa; sin ese espacio se ve como si
    perteneciera a la opción equivocada)."""
    return {nombre: (' ☒' if clave == valor_seleccionado else ' ☐') for clave, nombre in mapa_chk.items()}


def _celdas_unicas(documento):
    """Itera las celdas de todas las tablas del documento sin repetir. Una
    celda fusionada a lo ancho de varias columnas aparece una vez por cada
    columna que cubre en `fila.cells` -- se descartan esos duplicados. OJO:
    se descarta por el elemento `_tc` en sí, no por `id(_tc)`, porque lxml
    crea un proxy Python efímero en cada acceso y `id(_tc)` puede reciclarse
    entre columnas distintas."""
    vistas = set()
    for tabla in documento.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                if celda._tc in vistas:
                    continue
                vistas.add(celda._tc)
                yield celda


def _parrafos_del_documento(documento):
    """Todos los párrafos del documento: los del cuerpo principal y los de
    cada celda de cada tabla."""
    yield from documento.paragraphs
    for celda in _celdas_unicas(documento):
        yield from celda.paragraphs


def _insertar_imagen(parrafo, imagen):
    run = parrafo.add_run()
    run.add_picture(imagen, height=ALTO_FIRMA)


def _insertar_firma_antes_de(buffer_docx, imagen_firma, coincide):
    """Abre el .docx ya renderizado e inserta la imagen de la firma (PNG con
    fondo transparente) en un párrafo nuevo, justo encima del primer
    párrafo cuyo texto (en mayúsculas) cumple `coincide`."""
    documento = DocumentoWord(buffer_docx)
    for parrafo in _parrafos_del_documento(documento):
        if coincide(parrafo.text.strip().upper()):
            nuevo_parrafo = parrafo.insert_paragraph_before()
            nuevo_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _insertar_imagen(nuevo_parrafo, imagen_firma)
            break

    salida = BytesIO()
    documento.save(salida)
    salida.seek(0)
    return salida


def _insertar_firma_en_caja(buffer_docx, imagen_firma, texto_titulo):
    """Abre el .docx ya renderizado y coloca la imagen de la firma en el
    primer párrafo (vacío) de la celda reservada para firmar, justo debajo
    de la celda cuyo texto es `texto_titulo` (ej. "Firma del suscriptor"),
    dentro de la misma tabla."""
    documento = DocumentoWord(buffer_docx)
    for tabla in documento.tables:
        for indice_fila, fila in enumerate(tabla.rows):
            if fila.cells[0].text.strip().upper() == texto_titulo and indice_fila + 1 < len(tabla.rows):
                celda_destino = tabla.rows[indice_fila + 1].cells[0]
                parrafo_destino = celda_destino.paragraphs[0]
                parrafo_destino.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _insertar_imagen(parrafo_destino, imagen_firma)
                break

    salida = BytesIO()
    documento.save(salida)
    salida.seek(0)
    return salida


def generar_solicitud_credito(prestamo, imagen_firma=None):
    """Rellena la plantilla oficial de Solicitud de Crédito con los datos del
    préstamo y el expediente KYC del cliente. Devuelve un BytesIO listo para
    descargar (no se persiste en el modelo, se genera al vuelo).

    `imagen_firma` (BytesIO opcional) es la imagen de firma digital ya
    procesada (fondo blanco eliminado); si se recibe, se inserta sobre el
    texto "FIRMA DEL SOLICITANTE". Si es None, la solicitud se genera igual
    que antes, para que el firmante la imprima y la firme a mano."""
    if not PLANTILLA_SOLICITUD_PATH.exists():
        raise FileNotFoundError('No existe la plantilla de solicitud de crédito.')

    cliente = prestamo.cliente
    contexto = {
        # Datos del crédito
        'plazo': prestamo.plazo_meses,
        'monto_solicitado': f'L {prestamo.monto_solicitado:,.2f}',
        'fecha_solicitud': prestamo.fecha_solicitud.strftime('%d/%m/%Y'),
        **_marcar(_DESTINO_A_CHK, prestamo.destino),

        # 1. Identificación personal
        'primer_nombre': cliente.primer_nombre,
        'segundo_nombre': cliente.segundo_nombre,
        'primer_apellido': cliente.primer_apellido,
        'segundo_apellido': cliente.segundo_apellido,
        'fecha_nacimiento': cliente.fecha_nacimiento.strftime('%d/%m/%Y'),
        'lugar_nacimiento': cliente.lugar_nacimiento,
        'numero_identificacion': cliente.numero_identificacion,
        'nacionalidad': cliente.get_nacionalidad_display(),
        'profesion': cliente.profesion_ocupacion_oficio,
        **_marcar(_TIPO_ID_A_CHK, cliente.tipo_identificacion),
        **_marcar(_GENERO_A_CHK, cliente.genero),

        # 2. Contacto y domicilio
        'telefono_fijo': cliente.telefono_fijo,
        'celular': cliente.celular,
        'email_personal': cliente.email_personal,
        'direccion_domicilio': cliente.direccion_colonia_barrio,
        'calle_avenida': cliente.calle_avenida,
        'numero_casa': cliente.numero_casa,
        'punto_referencia': cliente.punto_referencia,
        'municipio': cliente.municipio,
        'departamento': cliente.departamento,
        'pais': cliente.pais,

        # 3. Estado civil y familia
        'numero_dependientes': cliente.numero_dependientes,
        'conyuge_nombre': cliente.nombre_conyuge,
        'conyuge_telefono_fijo': cliente.conyuge_telefono_fijo,
        'conyuge_celular': cliente.conyuge_celular,
        'conyuge_empresa': cliente.conyuge_empresa,
        **_marcar(_ESTADO_CIVIL_A_CHK, cliente.estado_civil),

        # 4. Referencias personales
        'referencia1_nombre': cliente.ref1_nombre,
        'referencia1_telefono_fijo': cliente.ref1_telefono_fijo,
        'referencia1_celular': cliente.ref1_celular,
        'referencia2_nombre': cliente.ref2_nombre,
        'referencia2_telefono_fijo': cliente.ref2_telefono_fijo,
        'referencia2_celular': cliente.ref2_celular,

        # 5. Información profesional y laboral
        'empresa_nombre': cliente.empresa_nombre,
        'fecha_ingreso': cliente.empresa_fecha_ingreso.strftime('%d/%m/%Y'),
        'anios_laborando': cliente.empresa_anios_laborando,
        'cargo_actual': cliente.cargo_actual,
        'empresa_telefono': cliente.empresa_telefono,
        'email_laboral': cliente.empresa_email,
        'empresa_direccion': cliente.empresa_direccion,
        'empresa_ciudad': cliente.empresa_ciudad,
        'empresa_municipio': cliente.empresa_municipio,
        'empresa_departamento': cliente.empresa_departamento,
        'nombre_gerente_rrhh': cliente.gerente_rrhh_nombre,
        'nombre_jefe_inmediato': cliente.jefe_inmediato_nombre,
        **_marcar(_TIPO_EMPRESA_A_CHK, cliente.tipo_empresa),
        **_marcar(_TIPO_EMPLEADO_A_CHK, cliente.tipo_empleado),

        # 6. Información financiera — el ingreso mensual ya no se recolecta en el
        # formulario, así que se deja en blanco en vez de imprimir un falso "L 0.00".
        'ingreso_mensual': '',
        **_marcar(_RANGO_SALARIO_A_CHK, cliente.rango_salario),

        # Ciudad donde se firma: siempre la sede de la institución, no depende
        # del cliente.
        'ciudad_firma': CIUDAD_INSTITUCION,

        **_CAMPOS_NO_RECOLECTADOS,
    }
    contexto = _con_sangria(contexto)

    documento = DocxTemplate(str(PLANTILLA_SOLICITUD_PATH))
    documento.render(contexto)

    buffer = BytesIO()
    documento.save(buffer)
    buffer.seek(0)

    if imagen_firma is not None:
        buffer = _insertar_firma_antes_de(buffer, imagen_firma, lambda texto: texto == TEXTO_FIRMA_SOLICITANTE)

    return buffer


def generar_recibo_desembolso(prestamo, imagen_firma=None):
    """Rellena la plantilla oficial de Recibo de Desembolso con los datos
    del préstamo ya desembolsado. Devuelve un BytesIO listo para descargar
    (no se persiste en el modelo, se genera al vuelo).

    `imagen_firma` (BytesIO opcional) es la imagen de firma digital ya
    procesada (fondo blanco eliminado); si se recibe, se inserta sobre la
    línea "FIRMA ___ HUELLA ___". Si es None, se genera igual que antes,
    para firmar a mano."""
    if not PLANTILLA_RECIBO_DESEMBOLSO_PATH.exists():
        raise FileNotFoundError('No existe la plantilla de recibo de desembolso.')
    if not prestamo.fecha_desembolso:
        raise ValueError('El préstamo todavía no ha sido desembolsado.')

    fecha = prestamo.fecha_desembolso
    contexto = {
        'cliente_nombre': prestamo.cliente.nombre_completo,
        'cliente_dni': prestamo.cliente.numero_identificacion,
        'monto_desembolsado': f'{prestamo.monto_solicitado:,.2f}',
        'dia_desembolso': fecha.day,
        'mes_desembolso': _MESES_ES[fecha.month],
        'anio_desembolso': fecha.year,
    }

    documento = DocxTemplate(str(PLANTILLA_RECIBO_DESEMBOLSO_PATH))
    documento.render(contexto)

    buffer = BytesIO()
    documento.save(buffer)
    buffer.seek(0)

    if imagen_firma is not None:
        # "FIRMA ___" y "HUELLA ___" viven en celdas separadas de una tabla
        # (sin bordes) para que la firma quede centrada solo sobre el ancho
        # de la columna de FIRMA, no sobre toda la línea (que también
        # abarcaría la columna de HUELLA).
        coincide = lambda texto: texto.startswith('FIRMA')
        buffer = _insertar_firma_antes_de(buffer, imagen_firma, coincide)

    return buffer


def generar_pagare(prestamo, imagen_firma=None):
    """Rellena la plantilla oficial de Pagaré con los datos del préstamo ya
    desembolsado. Devuelve un BytesIO listo para descargar (no se persiste
    en el modelo, se genera al vuelo).

    - "Fecha de expedición" es la fecha de desembolso (cuando se firma el
      pagaré, junto con el Recibo de Desembolso).
    - "Fecha de pago" es el vencimiento de la última cuota de la tabla de
      amortización (cuándo el préstamo queda totalmente pagado).
    - `imagen_firma` (BytesIO opcional) es la imagen de firma digital ya
      procesada; si se recibe, se coloca dentro de la caja "Firma del
      suscriptor". Si es None, se genera igual que antes, para firmar a mano.
    """
    if not PLANTILLA_PAGARE_PATH.exists():
        raise FileNotFoundError('No existe la plantilla de pagaré.')
    if not prestamo.fecha_desembolso:
        raise ValueError('El préstamo todavía no ha sido desembolsado.')

    ultima_cuota = prestamo.cuotas.order_by('-numero_cuota').first()
    if not ultima_cuota:
        raise ValueError('El préstamo no tiene tabla de amortización generada.')

    fecha_expedicion = prestamo.fecha_desembolso
    fecha_pago = ultima_cuota.fecha_vencimiento
    contexto = {
        'ciudad_pagare': CIUDAD_INSTITUCION,
        'dia_expedicion': fecha_expedicion.day,
        'mes_expedicion': _MESES_ES[fecha_expedicion.month],
        'anio_expedicion': fecha_expedicion.year,
        'lugar_pago': CIUDAD_INSTITUCION,
        'fecha_pago': f'{fecha_pago.day} de {_MESES_ES[fecha_pago.month]} de {fecha_pago.year}',
        'monto_en_letras': monto_en_letras(prestamo.monto_solicitado),
        'cliente_nombre': prestamo.cliente.nombre_completo,
        'cliente_direccion': prestamo.cliente.direccion_completa or 'No registrada',
    }

    documento = DocxTemplate(str(PLANTILLA_PAGARE_PATH))
    documento.render(contexto)

    buffer = BytesIO()
    documento.save(buffer)
    buffer.seek(0)

    if imagen_firma is not None:
        buffer = _insertar_firma_en_caja(buffer, imagen_firma, TEXTO_FIRMA_SUSCRIPTOR)

    return buffer


def generar_recibo_pago(transaccion):
    """Rellena la plantilla oficial de Recibo de Pago con los datos de un
    pago de cuota puntual (una TransaccionCaja de tipo PAGO_CUOTA). Devuelve
    un BytesIO listo para descargar (no se persiste en el modelo, se genera
    al vuelo). El "saldo pendiente" que se imprime es el saldo TOTAL del
    préstamo después de esa cuota (mismo campo que se ve en la tabla de
    amortización), no el saldo parcial de la cuota."""
    if not PLANTILLA_RECIBO_PAGO_PATH.exists():
        raise FileNotFoundError('No existe la plantilla de recibo de pago.')

    fecha = transaccion.fecha_hora
    contexto = {
        'cliente_nombre': transaccion.prestamo.cliente.nombre_completo,
        'cliente_dni': transaccion.prestamo.cliente.numero_identificacion,
        'monto_pagado': f'{transaccion.monto_pagado:,.2f}',
        'saldo_pendiente': f'{transaccion.cuota.saldo:,.2f}',
        'dia_pago': fecha.day,
        'mes_pago': _MESES_ES[fecha.month],
        'anio_pago': fecha.year,
    }

    documento = DocxTemplate(str(PLANTILLA_RECIBO_PAGO_PATH))
    documento.render(contexto)

    buffer = BytesIO()
    documento.save(buffer)
    buffer.seek(0)
    return buffer
