"""Genera la plantilla .docx base con placeholders Jinja (docxtpl) del
Pagaré. Reemplaza el archivo que se había subido como una imagen (una
captura de pantalla, sin texto editable) por un documento con el mismo
contenido pero como texto real, para poder rellenarlo con los datos del
préstamo igual que el resto de documentos. Se ejecuta una sola vez para
crear el archivo versionable en financiera/templates_docx/PAGARÉ.docx."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.core.management.base import BaseCommand
from docx.shared import Pt, RGBColor

PLANTILLA_PATH = Path(__file__).resolve().parent.parent.parent / 'templates_docx' / 'PAGARÉ.docx'

COLOR_CAPTION = RGBColor(0x6b, 0x72, 0x80)


def _agregar_valor(parrafo, texto):
    run = parrafo.add_run(texto)
    run.bold = True
    run.underline = True
    return run


def _agregar_caption(doc, texto):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(texto)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = COLOR_CAPTION
    return p


class Command(BaseCommand):
    help = 'Crea la plantilla .docx base para la generación de pagarés.'

    def handle(self, *args, **options):
        doc = Document()

        # --- Encabezado: título + lugar y fecha de expedición ---
        tabla_encabezado = doc.add_table(rows=1, cols=2)
        tabla_encabezado.style = 'Light Grid Accent 1'
        celda_titulo, celda_fecha = tabla_encabezado.rows[0].cells

        p_titulo = celda_titulo.paragraphs[0]
        p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_titulo = p_titulo.add_run('PAGARÉ')
        run_titulo.bold = True
        run_titulo.font.size = Pt(20)

        p_fecha = celda_fecha.paragraphs[0]
        p_fecha.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_fecha.add_run('En ')
        _agregar_valor(p_fecha, '{{ ciudad_pagare }}')
        p_fecha.add_run(' a ')
        _agregar_valor(p_fecha, '{{ dia_expedicion }}')
        p_fecha.add_run(' de ')
        _agregar_valor(p_fecha, '{{ mes_expedicion }}')
        p_fecha.add_run(' de ')
        _agregar_valor(p_fecha, '{{ anio_expedicion }}')
        p_fecha.add_run('.')

        p_fecha_caption = celda_fecha.add_paragraph()
        p_fecha_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_fecha_caption = p_fecha_caption.add_run('Lugar y fecha de expedición')
        run_fecha_caption.italic = True
        run_fecha_caption.font.size = Pt(9)
        run_fecha_caption.font.color.rgb = COLOR_CAPTION

        doc.add_paragraph()

        # --- Cuerpo: obligación de pago ---
        p1 = doc.add_paragraph()
        p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p1.add_run('Debo y pagaré incondicionalmente por este pagaré a la orden de ')
        _agregar_valor(p1, 'FINCEL')
        _agregar_caption(doc, 'Nombre de la persona a quien ha de pagarse')

        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p2.add_run('en ')
        _agregar_valor(p2, '{{ lugar_pago }}')
        p2.add_run(' el día ')
        _agregar_valor(p2, '{{ fecha_pago }}')
        _agregar_caption(doc, 'Lugar de pago / Fecha de pago')

        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p3.add_run('la cantidad de ')
        _agregar_valor(p3, '{{ monto_en_letras }}')

        doc.add_paragraph()
        doc.add_paragraph()

        # --- Datos del suscriptor (deudor) ---
        p_titulo_suscriptor = doc.add_paragraph()
        run_titulo_suscriptor = p_titulo_suscriptor.add_run('Datos del suscriptor')
        run_titulo_suscriptor.bold = True

        p_nombre = doc.add_paragraph()
        p_nombre.add_run('Nombre ')
        _agregar_valor(p_nombre, '{{ cliente_nombre }}')

        p_direccion = doc.add_paragraph()
        p_direccion.add_run('Dirección ')
        _agregar_valor(p_direccion, '{{ cliente_direccion }}')

        doc.add_paragraph()

        # --- Firma del suscriptor (queda en blanco, se firma a mano) ---
        tabla_firma = doc.add_table(rows=2, cols=1)
        tabla_firma.style = 'Light Grid Accent 1'
        celda_firma_titulo, celda_firma_espacio = tabla_firma.rows[0].cells[0], tabla_firma.rows[1].cells[0]
        p_firma_titulo = celda_firma_titulo.paragraphs[0]
        p_firma_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_firma_titulo = p_firma_titulo.add_run('Firma del suscriptor')
        run_firma_titulo.bold = True
        celda_firma_espacio.add_paragraph()
        celda_firma_espacio.add_paragraph()

        PLANTILLA_PATH.parent.mkdir(parents=True, exist_ok=True)
        doc.save(PLANTILLA_PATH)

        self.stdout.write(self.style.SUCCESS(f'Plantilla creada en {PLANTILLA_PATH}'))
