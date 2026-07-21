"""Genera la plantilla .docx base con placeholders Jinja (docxtpl) que usa
el servicio de generación de contratos. Se ejecuta una sola vez para crear
el archivo versionable en financiera/templates_docx/contrato_credito.docx."""
from pathlib import Path

from django.core.management.base import BaseCommand
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

PLANTILLA_PATH = Path(__file__).resolve().parent.parent.parent / 'templates_docx' / 'contrato_credito.docx'


class Command(BaseCommand):
    help = 'Crea la plantilla .docx base para la generación de contratos de crédito.'

    def handle(self, *args, **options):
        doc = Document()

        titulo = doc.add_heading('CONTRATO DE CRÉDITO', level=1)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitulo = doc.add_paragraph('MicroFinanzas Pro')
        subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitulo.runs[0].bold = True

        doc.add_paragraph('Código de crédito: {{ codigo_credito }}')
        doc.add_paragraph('Fecha de aprobación: {{ fecha_aprobacion }}')

        doc.add_heading('1. Partes', level=2)
        doc.add_paragraph(
            'Entre MicroFinanzas Pro, en adelante "LA INSTITUCIÓN", y {{ cliente_nombre }}, '
            'identificado(a) con DNI {{ cliente_dni }}, domiciliado(a) en {{ cliente_direccion }}, '
            'en adelante "EL CLIENTE", se celebra el presente contrato de crédito bajo las '
            'siguientes cláusulas.'
        )

        doc.add_heading('2. Condiciones del Crédito', level=2)
        tabla_condiciones = doc.add_table(rows=5, cols=2)
        tabla_condiciones.style = 'Light Grid Accent 1'
        filas = [
            ('Monto solicitado', '{{ monto_solicitado }}'),
            ('Tasa de interés anual', '{{ tasa_interes_anual }}%'),
            ('Plazo', '{{ plazo_meses }} meses'),
            ('Frecuencia de pago', '{{ frecuencia_pago }}'),
            ('Número de cuotas', '{{ numero_cuotas }}'),
        ]
        for fila, (etiqueta, valor) in zip(tabla_condiciones.rows, filas):
            fila.cells[0].text = etiqueta
            fila.cells[1].text = valor

        doc.add_heading('3. Cláusulas', level=2)
        doc.add_paragraph(
            'EL CLIENTE se compromete a pagar el monto total del crédito más los intereses '
            'pactados de acuerdo con la Tabla de Amortización adjunta en la cláusula 4, en las '
            'fechas de vencimiento indicadas. El incumplimiento de pago genera el estado de mora '
            'conforme a las políticas de LA INSTITUCIÓN.'
        )

        doc.add_heading('4. Tabla de Amortización', level=2)
        # docxtpl reemplaza la fila COMPLETA que contiene un tag {%tr ...%},
        # por eso las filas de apertura/cierre del bucle van separadas de la
        # fila con los datos de cada cuota.
        tabla_cuotas = doc.add_table(rows=4, cols=5)
        tabla_cuotas.style = 'Light Grid Accent 1'

        encabezados = tabla_cuotas.rows[0].cells
        for celda, texto in zip(encabezados, ['N°', 'Vencimiento', 'Capital', 'Interés', 'Total Cuota']):
            celda.text = texto
            celda.paragraphs[0].runs[0].bold = True

        tabla_cuotas.rows[1].cells[0].text = '{%tr for cuota in cuotas %}'

        fila_datos = tabla_cuotas.rows[2].cells
        valores_fila = [
            '{{ cuota.numero_cuota }}',
            '{{ cuota.fecha_vencimiento }}',
            '{{ cuota.monto_capital }}',
            '{{ cuota.monto_interes }}',
            '{{ cuota.monto_total_cuota }}',
        ]
        for celda, texto in zip(fila_datos, valores_fila):
            celda.text = texto

        tabla_cuotas.rows[3].cells[0].text = '{%tr endfor %}'

        doc.add_heading('5. Firmas', level=2)
        firmas = doc.add_paragraph()
        firmas.add_run('\n\n_____________________________\t\t_____________________________\n')
        firmas.add_run('Por LA INSTITUCIÓN\t\t\t\tEL CLIENTE\n')
        firmas.add_run('{{ cliente_nombre }}')

        for parrafo in [subtitulo]:
            for run in parrafo.runs:
                run.font.size = Pt(12)

        PLANTILLA_PATH.parent.mkdir(parents=True, exist_ok=True)
        doc.save(PLANTILLA_PATH)

        self.stdout.write(self.style.SUCCESS(f'Plantilla creada en {PLANTILLA_PATH}'))
